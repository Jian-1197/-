
# 这些包是处理文件用的
import re

import pandas as pd
from openpyxl import Workbook   
from openpyxl.utils.dataframe import dataframe_to_rows  
from openpyxl.styles import Alignment, Font, Border, Side   
from openpyxl.utils import get_column_letter

from docx import Document 
from docx.shared import Pt, Inches 
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT   # 段落对齐
from docx.oxml.ns import qn                         # 中文格式
import subprocess                                   # 用命令行实现docx-to-pdf

# 样式设置变量 (脚本一)
TITLE_FONT_1 = Font(size=18, bold=True)                                       # 标题加粗, 18号字体
DATE_FONT_1 = Font(size=14, bold=True)                                        # 日期加粗, 14号字体
HEADER_FONT_1 = Font(size=12, bold=True)                                      # 表头加粗, 12号字体
CONTENT_FONT_1 = Font(size=12)                                                # 其他内容, 12号字体    
BORDER_STYLE_1 = Border(left=Side(style="thin"), right=Side(style="thin"),    # 边框样式
                      top=Side(style="thin"), bottom=Side(style="thin"))    
CENTER_ALIGNMENT_1 = Alignment(horizontal="center", vertical="center")        # 居中对齐

def set_style(ws, start_row, end_row, start_col, end_col, alignment=None, font=None, border=None, col_width=20, row_height=27):
    """为指定区域的单元格设置边框、对齐方式、字体、列宽和行高"""
    for row in ws.iter_rows(min_row=start_row, max_row=end_row, min_col=start_col, max_col=end_col):
        for cell in row:
            if border:
                cell.border = border
            if alignment:
                cell.alignment = alignment
            if font:
                cell.font = font
    
    for col_num in range(start_col, end_col + 1):
        ws.column_dimensions[get_column_letter(col_num)].width = col_width
    for row_num in range(start_row, end_row + 1):
        ws.row_dimensions[row_num].height = row_height

def style_sheet(ws,alignment,font_head,font_content,border,row_height=27) :
    # 设置行高和单元格样式
    for row in ws.iter_rows(min_row=2):  # 跳过表头
        ws.row_dimensions[row[0].row].height = row_height
        for cell in row:
            cell.alignment = alignment
            cell.font = font_content
            cell.border = border
    
    # 设置表头字体样式
    for row in ws.iter_rows(min_row=1, max_row=1):
        ws.row_dimensions[row[0].row].height = 27
        for cell in row:
            cell.font = font_head
            cell.alignment = alignment
            cell.border = border

def process_confirm_sheets(data,study_year,semester,start_year,start_month,start_day,end_year,end_month,end_day):

    def is_chinese(name):
        return bool(re.search(r'^[\u4e00-\u9fa5]+$', name))
    data["姓名是否中文"] = data["姓名"].apply(is_chinese)    

    # 获取不重复班级名称,不包括留学生班级以及21级以下的班级
    unique_classes = [
        cls for cls in data["班级"].dropna().unique()
        if not re.search(r'全英文班|2001|2002|1901|1902|1801|1802', cls)
    ]

    
    # 逐个班级处理数据并生成新表格
    def create_confirm_sheet():
        for class_name in unique_classes:
            # 筛选班级数据
            class_data = data[data["班级"] == class_name]

            # 过滤条件: 旷课课时大于0
            filtered_data = class_data[(class_data["旷课课时"] > 0)]

            # 选择所需列并添加"确认签字"列
            filtered_data = filtered_data[["姓名", "旷课次数", "迟到次数", "早退次数", "旷课课时"]]
            filtered_data["确认签字"] = ""

            # 按旷课课时降序, 其次按旷课次数降序排序
            filtered_data = filtered_data.sort_values(by=["旷课课时", "旷课次数"], ascending=[False, False])

            # 创建新的工作簿
            wb = Workbook()
            ws = wb.active

            # 添加标题行并合并单元格
            ws.merge_cells("A1:F1")
            ws["A1"] = f"{study_year}年度{class_name}上课啦扣分情况统计表"
            set_style(ws, 1, 1, 1, 6, CENTER_ALIGNMENT_1, TITLE_FONT_1, BORDER_STYLE_1, col_width=30, row_height=30)

            # 添加日期行并合并单元格
            ws.merge_cells("A2:F2")
            ws["A2"] = f"{start_year}年{start_month}月{start_day}日-{end_year}年{end_month}月{end_day}日"
            set_style(ws, 2, 2, 1, 6, CENTER_ALIGNMENT_1, DATE_FONT_1, BORDER_STYLE_1, col_width=30, row_height=25)

            # 添加表头
            header = ["姓名", "旷课次数", "迟到次数", "早退次数", "旷课课时", "确认签字"]
            ws.append(header)
            set_style(ws, 3, 3, 1, 6, CENTER_ALIGNMENT_1, HEADER_FONT_1, BORDER_STYLE_1)

            # 写入数据
            for row in dataframe_to_rows(filtered_data, index=False, header=False):
                ws.append(row)

            # 设置数据行样式
            set_style(ws, 4, ws.max_row, 1, 6, CENTER_ALIGNMENT_1, CONTENT_FONT_1, BORDER_STYLE_1)

            # 保存文件
            output_file = f"{output_folder_1}/{class_name}.xlsx"
            wb.save(output_file)

    def create_summary_sheet():
        filtered_data = data[(data["旷课课时"] > 0) & data["姓名是否中文"]]
        filtered_data = filtered_data[["学号", "姓名", "班级","旷课次数", "迟到次数", "早退次数", "旷课课时"]]
        filtered_data["学号"] = filtered_data["学号"].astype(str)
        filtered_data = filtered_data.sort_values(by=["旷课课时", "旷课次数"], ascending=[False, False])

        wb = Workbook()
        ws = wb.active

        colum_widths = [15, 11, 35, 10, 10, 10, 10]
        header = ["学号", "姓名", "班级", "旷课次数", "迟到次数", "早退次数", "旷课课时"]
        ws.append(header)

        # 设置列宽
        for idx, col_width in enumerate(colum_widths, start=1):
            ws.column_dimensions[get_column_letter(idx)].width = col_width

        for row in dataframe_to_rows(filtered_data, index=False, header=False):
            ws.append(row)
        style_sheet(ws,CENTER_ALIGNMENT_1,HEADER_FONT_1,CONTENT_FONT_1,BORDER_STYLE_1)

        output_file = f"{output_folder_2}/{study_year}上课啦汇总.xlsx"
        wb.save(output_file)
    
    def create_circular_sheet():
        # 创建一个新的Excel工作簿
        wb = Workbook()

        # 定义表头
        header_common = ["序号", "姓名", "时间", "班级", "原因",""]
        header_sheet2 = ["序号", "姓名", "时间", "班级", "原因","","类型"]

        # 定义相关内容
        title = f"{study_year}学年{semester}计算机科学与技术学院违规违纪名单"
        time_str = f"{study_year}学年{semester}学期"
        reason_1 = "旷课课时满5学时不足10学时"
        reason_2_1 = "旷课课时满10学时不足20学时"
        reason_2_2 = "旷课课时满20学时不足30学时"
        reason_3 = "旷课课时满30学时不足40学时"
        reason_4 = "旷课课时40学时以上"

        # 定义列宽
        column_widths_common = [5, 12, 35, 40, 31, 6]
        column_widths_sheet2 = [5, 12, 35, 40, 31, 6, 10]

        # 定义样式
        center_alignment = Alignment(horizontal='center', vertical='center')
        title_font = Font(size=16,bold=True)
        header_font = Font(size=12,bold=True)
        content_font = Font(size=12)
        border_style = Border(left=Side(style='thin'), 
                            right=Side(style='thin'), 
                            top=Side(style='thin'), 
                            bottom=Side(style='thin'))
        # 写表函数
        def write_sheet(ws, data, headers, col_widths, title, time_str, reason_str=None, type_col_name=None):
            # 合并第一行并写入标题
            ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
            ws.cell(row=1, column=1).value = title

            # 写入表头
            ws.append(headers)
            ws.merge_cells(start_row=2, start_column=5, end_row=2, end_column=6)

            # 设置列宽
            for idx, col_width in enumerate(col_widths, start=1):
                ws.column_dimensions[get_column_letter(idx)].width = col_width

            # 写入数据
            for index, (_, row) in enumerate(data.iterrows()):
                if type_col_name:
                    if 10 <= row['旷课课时'] < 20:
                        type_value = "警告"
                        reason_str = reason_2_1
                    elif 20 <= row['旷课课时'] < 30:
                        type_value = "严重警告"
                        reason_str = reason_2_2
                    else:
                        type_value = ""
                    ws.append([index + 1, row['姓名'], time_str, row['班级'], reason_str, row['旷课课时'], type_value])
                else:
                    ws.append([index + 1, row['姓名'], time_str, row['班级'], reason_str, row['旷课课时']])

            # 设置样式
            # 标题
            for row in ws.iter_rows(min_row=1, max_row=1):
                ws.row_dimensions[row[0].row].height = 27
                for cell in row:
                    cell.alignment = center_alignment
                    cell.font = title_font
                    cell.border = border_style
            # 表头
            for row in ws.iter_rows(min_row=2, max_row=2):
                ws.row_dimensions[row[0].row].height = 27
                for cell in row:
                    cell.font = header_font
                    cell.alignment = center_alignment
                    cell.border = border_style
            # 数据
            for row in ws.iter_rows(min_row=3):
                ws.row_dimensions[row[0].row].height = 27
                for cell in row:
                    cell.alignment = center_alignment
                    cell.font = content_font
                    cell.border = border_style

        # 筛选数据
        filtered_data = data[(data["旷课课时"] >=5) & data["姓名是否中文"]]
        filtered_data = filtered_data.sort_values(by=["旷课课时", "旷课次数"], ascending=[True,True])
        sheet1_data = filtered_data[(filtered_data['旷课课时'] >= 5) & (filtered_data['旷课课时'] < 10)]
        sheet2_data = filtered_data[(filtered_data['旷课课时'] >= 10) & (filtered_data['旷课课时'] < 30)]
        sheet3_data = filtered_data[(filtered_data['旷课课时'] >= 30) & (filtered_data['旷课课时'] < 40)]
        sheet4_data = filtered_data[filtered_data['旷课课时'] >= 40]

        # 写入Sheet1
        sheet1 = wb.active
        sheet1.title = "sheet1"
        write_sheet(sheet1, sheet1_data, header_common, column_widths_common, title, time_str, reason_1)

        # 写入Sheet2
        sheet2 = wb.create_sheet(title="sheet2")
        write_sheet(sheet2, sheet2_data, header_sheet2, column_widths_sheet2,title, time_str, type_col_name="类型")

        # 写入Sheet3
        sheet3 = wb.create_sheet(title="sheet3")
        write_sheet(sheet3, sheet3_data, header_common, column_widths_common, title, time_str, reason_3)

        # 写入Sheet4
        sheet4 = wb.create_sheet(title="sheet4")
        write_sheet(sheet4, sheet4_data, header_common, column_widths_common, title, time_str, reason_4)

        # 保存文件
        output_file = f"{output_folder_2}/{study_year}上课啦违规违纪名单.xlsx"
        wb.save(output_file)
    
    create_confirm_sheet()
    create_summary_sheet()
    create_circular_sheet()
    

# 样式设置变量 (脚本二)
HEADER_FONT_2 = Font(size=11, bold=True)  # 表头字体样式
CONTENT_FONT_2 = Font(size=11)            # 内容字体样式
BORDER_STYLE_2 = Border(                  # 单元格边框样式    
    left=Side(style="thin"),
    right=Side(style="thin"),
    top=Side(style="thin"),
    bottom=Side(style="thin")
)
CENTER_ALIGNMENT_2 = Alignment(horizontal="center", vertical="center")  # 居中对齐

def process_attendance_files(data, date, year, month, day):

    # 检查姓名列是否符合条件(中文或非中文)  
    def is_chinese(name):
        return bool(re.search(r'^[\u4e00-\u9fa5]+$', name))

    data["姓名是否中文"] = data["姓名"].apply(is_chinese)

    # 创建 Excel 表格
    def create_excel():
        wb = Workbook()
        
        # 本科生 sheet
        benke = wb.active
        benke.title = "本科生"
        cols = ["学号", "姓名", "学院", "班级", "旷课次数", "迟到次数", "早退次数", "旷课课时"]
        filtered_data = data[(data["旷课课时"] >= 2) & data["姓名是否中文"]]
        filtered_data = filtered_data.sort_values(by=["旷课课时", "旷课次数"], ascending=[False, False])
        write_sheet(benke, filtered_data, cols, [15, 10, 10, 10, 10, 10, 10, 10], ["学院", "班级"])
        
        # 留学生 sheet
        international = wb.create_sheet(title="留学生")
        filtered_data = data[(data["旷课课时"] >= 2) & ~data["姓名是否中文"]]
        filtered_data = filtered_data.sort_values(by=["旷课课时", "旷课次数"], ascending=[False, False])
        write_sheet(international, filtered_data, cols, [15, 35, 10, 10, 10, 10, 10, 10], ["学院", "班级"])

        # 保存文件
        excel_output = f"{output_folder_3}/计算机科学与技术学院学生第{date}上课啦系统缺勤情况.xlsx"
        wb.save(excel_output)

    # 写入 Excel Sheet
    def write_sheet(ws, data, columns, col_widths, hidden_cols):
        # 将学号列转换为字符串格式
        data["学号"] = data["学号"].astype(str)
        
        # 选择所需列并添加表头
        data = data[columns]
        ws.append(columns)
        
        # 设置列宽
        for idx, col_width in enumerate(col_widths, start=1):
            ws.column_dimensions[get_column_letter(idx)].width = col_width
        
        # 写入数据
        for _, row in data.iterrows():
            ws.append(row.tolist())
        
        # 隐藏指定列
        for col in hidden_cols:
            col_idx = columns.index(col) + 1
            ws.column_dimensions[get_column_letter(col_idx)].hidden = True
        
        # 设置样式
        style_sheet(ws)

    # 样式设置
    def style_sheet(ws):
        # 设置行高和单元格样式
        for row in ws.iter_rows(min_row=2):  # 跳过表头
            ws.row_dimensions[row[0].row].height = 27
            for cell in row:
                cell.alignment = CENTER_ALIGNMENT_2
                cell.font = CONTENT_FONT_2
                cell.border = BORDER_STYLE_2
        
        # 设置表头字体样式
        for row in ws.iter_rows(min_row=1, max_row=1):
            ws.row_dimensions[row[0].row].height = 27
            for cell in row:
                cell.font = HEADER_FONT_2
                cell.alignment = CENTER_ALIGNMENT_2
                cell.border = BORDER_STYLE_2


    def create_docx():
        # 创建 Word 文档
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'  # 必须先设置font.name
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        title = f"第{date}“上课啦”考勤通报"
        para_1 = doc.add_paragraph()
        run_1 = para_1.add_run(title)
        run_1.font.size = Pt(22)
        run_1.font.bold = True
        para_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 设置标题居中

        # 段落内容
        content = f"以下同学在第{date}“上课啦”考勤中未正常出勤, 旷课学时计入个人档案, 并纳入日常考评。"
        para_2 = doc.add_paragraph()
        run_2 = para_2.add_run(content)
        run_2.font.size = Pt(16)
        para_2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 设置内容两端对齐

        # 添加表格
        filtered_data = data[(data["旷课课时"] >= 2) & data["姓名是否中文"]]
        filtered_data = filtered_data.sort_values(by=["旷课课时", "旷课次数"], ascending=[False, False])

        table = doc.add_table(rows=1, cols=6, style="Table Grid")
        col_width_dict = {0: 1.6, 1: 1.12, 2: 0.7638, 3: 0.7638, 4: 0.7638, 5: 0.7638}
        row_height = Pt(25)
        
        # 设置列宽
        for col_num in range(6):
            table.cell(0, col_num).width = Inches(col_width_dict[col_num])

        # 设置表头
        headers = ["学号", "姓名", "旷课次数", "迟到次数", "早退次数", "旷课课时"]
        header_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            cell = header_cells[idx]
            cell.text = header                      # 设置单元格文本
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True            # 设置文字加粗

        # 添加数据行
        for _, row in filtered_data.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row["学号"])
            row_cells[1].text = row["姓名"]
            row_cells[2].text = str(row["旷课次数"])
            row_cells[3].text = str(row["迟到次数"])
            row_cells[4].text = str(row["早退次数"])
            row_cells[5].text = str(row["旷课课时"])

        # 全局设置行高和垂直对齐方式
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = 1         # 单元格垂直居中
                row.height = row_height             # 设置行高
                for paragraph in cell.paragraphs:
                    paragraph.alignment = 1         # 设置单元格文本居中
                    for run in paragraph.runs:
                        run.font.size = Pt(11)

        # 添加说明
        note_1 = "\n注:\n一、根据学生手册中“浙江师范大学学生违纪处分规定”中第三章第二十七条规定,学生一学期内旷课累计满10学时的, 给予警告处分; 满20学时的, 给予严重警告处分;满30学时的, 给与记过处分; 满40学时的, 给与留校察看处分; 因旷课屡次受到纪律处分并经教育不改的, 可以给予开除学籍处分;\n二、学时计算方法如下: \n(1) 旷课1小节为1学时, 未经请假缺勤1天, 不足5学时按5学时计; 超过5学时的, 按实际学时数计; \n(2) 学生无故迟到或早退达3次, 作旷课1学时计; "

        para_3 = doc.add_paragraph()
        run_3 = para_3.add_run(note_1)
        run_3.font.size = Pt(12)
        run_3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 设置内容两端对齐


        note_3 = "特此通报！"
        para_4 = doc.add_paragraph()
        run_4 = para_4.add_run(note_3)
        run_4.font.size = Pt(16)
        run_4.font.bold = True
        run_4.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 设置内容左对齐

        para_5 = doc.add_paragraph()
        para_5.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 设置段落右对齐
        run_5 = para_5.add_run(f"计算机科学与技术学院学工办\n{year}年{month}月{day}日")
        run_5.font.size = Pt(16)


        # 保存文件
        docx_output = f"{output_folder_3}/计算机科学与技术学院学生第{date}上课啦系统缺勤通报.docx"
        doc.save(docx_output)
        return docx_output

    def convert_docx_to_pdf(input_file, output_format='pdf',output_dir='.'):
        
        
        # 构建命令字符串，包含输出目录参数
        command = [
            'soffice',  # LibreOffice/OpenOffice 的命令行工具
            '--headless',  # 不显示图形用户界面
            '--invisible',  # 运行时不可见
            '--convert-to', output_format,  # 转换格式为目标格式
            '--outdir', output_dir,  # 指定输出目录
            input_file  # 输入文件路径
        ]
        
        try:
            # 使用 subprocess.run 来执行命令
            result = subprocess.run(command, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            print("转换成功:", result.stdout.decode())
        except subprocess.CalledProcessError as e:
            print("转换失败:", e.stderr.decode())



    # 执行
    create_excel()
    docx_output = create_docx()
    convert_docx_to_pdf(docx_output,output_dir=f'{output_folder_3}')

# 这些包是压缩和删除文件用的
import os
import zipfile
import shutil
import atexit
import streamlit as st
import warnings
warnings.filterwarnings("ignore") # 要是要查bug或是优化代码，把这行注释掉

def zip_files(dirs, output):
    # 创建压缩文件
    zip_output = f"{output}.zip"
    with zipfile.ZipFile(zip_output, "w") as zipf:
        for dir in dirs:
            for root, _, files in os.walk(dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    if len(dirs) > 1:
                        arcname = os.path.join(os.path.basename(dir), os.path.relpath(file_path, start=dir))
                    else:
                        arcname = os.path.relpath(file_path, start=dir)
                    zipf.write(file_path, arcname)
    zipf.close()

def delete_files_and_folders(paths):
    for path in paths:
        if os.path.exists(path):
            if os.path.isdir(path):
                shutil.rmtree(path)
            else:
                os.remove(path)
def cleanup():
    delete_files_and_folders(["上课啦确认签字表", "学期汇总表+违规违纪名单"])

atexit.register(cleanup)

st.title("上课啦表格制作工具")

tab1, tab2 = st.tabs(["确认签字表", "考勤表"])

with tab1:
    st.header("上传并处理学期考勤数据")
    uploaded_file_attendance = st.file_uploader("上传学期考勤数据", type=["xls", "xlsx"])
    study_year_options = ["2024-2025", "2025-2026", "2026-2027", "2027-2028"]
    study_year = st.selectbox("选择学年", study_year_options)
    custom_study_year = st.text_input("或手动输入学年,但是注意格式", "")

    if custom_study_year:
        study_year = custom_study_year

    semester_options = ["第一学期", "第二学期"]
    semester = st.selectbox("学期", semester_options)
    start_date = st.date_input("学期开始日期")
    end_date = st.date_input("学期结束日期")
    if start_date and end_date:
        start_year = start_date.strftime('%Y')
        start_month = start_date.strftime('%m')
        start_day = start_date.strftime('%d')
        end_year = end_date.strftime('%Y')
        end_month = end_date.strftime('%m')
        end_day = end_date.strftime('%d')
    if uploaded_file_attendance is not None:

        global output_folder_1, output_folder_2

        output_folder_1 = "上课啦确认签字表"
        if not os.path.exists(output_folder_1):
            os.makedirs(output_folder_1)

        output_folder_2 = f"学期汇总表+违规违纪名单"
        if not os.path.exists(output_folder_2):
            os.makedirs(output_folder_2)

        data_attendance = pd.read_excel(uploaded_file_attendance)
        st.write("Data Preview:")
        st.dataframe(data_attendance.head())
        
        flag_1 = False

        if st.button("生成确认签字表"):
            with st.spinner('生成中...'):
                process_confirm_sheets(data_attendance, study_year=study_year, 
                semester=semester, start_year=start_year, start_month=start_month, 
                start_day=start_day, end_year=end_year, end_month=end_month, end_day=end_day)

                flag_1 = True
            st.success("生成成功！")
        if flag_1:
            zip_files([output_folder_1, output_folder_2], "确认签字表+汇总表")
            with open(f'确认签字表+汇总表.zip', 'rb') as f:
                with st.spinner('下载中...'):
                    st.download_button(
                        label='下载确认签字表+汇总表',
                        data=f,
                        file_name='确认签字表+汇总表.zip',
                        mime='application/zip',
                        on_click=lambda:delete_files_and_folders([output_folder_1, output_folder_2, '确认签字表+汇总表.zip'])
                    )

with tab2:
    st.header("上传并处理周/月考勤数据")
    uploaded_file_1 = st.file_uploader("上传周/月考勤数据", type=["xls", "xlsx"])
    uploaded_file_2 = st.file_uploader("上传考勤明细表", type=["xls", "xlsx"])
    selected_date = ['一周','二周','三周','四周','五周','六周','七周','八周','九周','十周',
                     '十一周','十二周','十三周','十四周','十五周','十六周','十七周','十八周',
                     '一月','二月','三月','四月','五月','六月','七月','八月','九月','十月',
                     '十一月','十二月']
    date = st.selectbox("选择第几周/月", selected_date)
    calendar = st.date_input("选择做表日期")
    if calendar:
        year = calendar.strftime('%Y')
        month = calendar.strftime('%m')
        day = calendar.strftime('%d')
    if uploaded_file_1 is not None and uploaded_file_2 is not None:
        data = pd.read_excel(uploaded_file_1)
        data_1 = pd.read_excel(uploaded_file_2)
        
        st.write("Data Preview:")
        st.dataframe(data.head())

        flag_2 = False
        if st.button("生成考勤通报"):
            
            with st.spinner('保存并生成文件中...'):
                global output_folder_3
                output_folder_3 = f"第{date}"
                if not os.path.exists(output_folder_3):
                    os.makedirs(output_folder_3)
                # 文件1重命名并保存
                file_name_1 = f"原始数据.xlsx"
                file_path_1 = os.path.join(output_folder_3, file_name_1)
                data["学号"] = data["学号"].astype(str)
                data.to_excel(file_path_1, index=False)

                # 文件2重命名并保存
                file_name_2 = f"计算机科学与技术学院第{date}上课啦考勤明细.xlsx"
                file_path_2 = os.path.join(output_folder_3, file_name_2)
                data_1["学号"] = data_1["学号"].astype(str)
                data_1["课程编号"] = data_1["课程编号"].astype(str)
                data_1.to_excel(file_path_2, index=False)

                process_attendance_files(data, date, year, month, day)
                flag_2 = True

                st.success(f"上传文件已保存，并且考勤通报生成成功！")
        if flag_2:
            zip_files([output_folder_3], f"{output_folder_3}")
            with open(f'{output_folder_3}.zip', 'rb') as f:
                with st.spinner('下载中...'):
                    st.download_button(
                        label='下载考勤通报',
                        data=f,
                        file_name=f'{output_folder_3}.zip',
                        mime='application/zip',
                        on_click=lambda:delete_files_and_folders([output_folder_3, f'{output_folder_3}.zip'])
                    )
